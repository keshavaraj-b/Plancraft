"""
Export Tools — export generated stories to Jira or Word format.
"""

import json
from datetime import datetime

from src.app import mcp
from src.state import state


@mcp.tool()
async def export_jira(project_key: str, dry_run: bool = True) -> dict:
    """
    Export generated stories to Jira (placeholder — returns preview only).

    Args:
        project_key: Jira project key (e.g., "CITI", "PASS")
        dry_run: If True, returns preview without pushing. Always True for now.

    Returns:
        Preview of what would be pushed to Jira.
    """
    if not state.last_stories:
        return {"error": "No stories generated yet. Run generate_stories_context first."}

    # Build Jira-formatted preview
    jira_items = []
    for epic in state.last_stories:
        epic_item = {
            "type": "Epic",
            "project": project_key,
            "summary": epic.get("title", "Untitled Epic"),
            "description": epic.get("description", ""),
            "children": [],
        }
        for feature in epic.get("features", []):
            for story in feature.get("user_stories", []):
                story_item = {
                    "type": "Story",
                    "project": project_key,
                    "summary": story.get("title", ""),
                    "description": story.get("description", ""),
                    "story_points": story.get("story_points", 0),
                    "priority": story.get("priority", "Medium"),
                    "acceptance_criteria": story.get("acceptance_criteria", []),
                    "subtasks": [{"type": "Sub-task", "summary": t} for t in story.get("tasks", [])],
                }
                epic_item["children"].append(story_item)
        jira_items.append(epic_item)

    return {
        "dry_run": dry_run,
        "project_key": project_key,
        "total_items": sum(1 + len(e["children"]) for e in jira_items),
        "preview": jira_items,
        "message": "Jira push is not yet implemented. This is a preview of what would be created."
        if dry_run
        else "NOT IMPLEMENTED — set dry_run=True",
    }


@mcp.tool()
async def export_word(stories_json: str, output_path: str | None = None) -> dict:
    """
    Export generated stories to a formatted Word (.docx) document.

    Creates a professional Word document with epics, features, user stories,
    acceptance criteria, tasks, and a summary table.

    Args:
        stories_json: JSON string of generated stories (epics/features/stories structure).
        output_path: File path to save the .docx file. If None, saves to output/stories_<timestamp>.docx

    Returns:
        Dictionary with file path of the generated document and summary.
    """
    from pathlib import Path

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    try:
        stories = json.loads(stories_json)
    except json.JSONDecodeError as e:
        return {"error": f"Invalid JSON: {e}"}

    epics = stories.get("epics", [])
    if not epics:
        return {"error": "No epics found in the stories JSON."}

    doc = Document()

    # Title page
    title = doc.add_heading("Development Plan — User Stories", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    client = state.connected_client or "Unknown Client"
    doc.add_paragraph(f"Client: {client}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Total Story Points: {stories.get('total_story_points', 'N/A')}")

    # Warnings section
    warnings = stories.get("warnings", [])
    if warnings:
        doc.add_heading("Warnings", level=1)
        for w in warnings:
            p = doc.add_paragraph(w, style="List Bullet")
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # Assumptions section
    assumptions = stories.get("assumptions", [])
    if assumptions:
        doc.add_heading("Assumptions", level=1)
        for a in assumptions:
            doc.add_paragraph(a, style="List Bullet")

    doc.add_page_break()

    # Epics, Features, Stories
    total_stories = 0
    for epic_idx, epic in enumerate(epics, 1):
        doc.add_heading(f"Epic {epic_idx}: {epic.get('title', 'Untitled')}", level=1)
        if epic.get("description"):
            doc.add_paragraph(epic["description"])

        for feature in epic.get("features", []):
            doc.add_heading(f"Feature: {feature.get('title', 'Untitled')}", level=2)

            for story in feature.get("user_stories", []):
                total_stories += 1
                _add_story_to_doc(doc, story, total_stories)

    # Summary table
    doc.add_page_break()
    doc.add_heading("Summary", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Epic"
    hdr[1].text = "Features"
    hdr[2].text = "Stories"
    hdr[3].text = "Story Points"

    for epic in epics:
        row = table.add_row().cells
        row[0].text = epic.get("title", "")
        features = epic.get("features", [])
        row[1].text = str(len(features))
        story_count = sum(len(f.get("user_stories", [])) for f in features)
        row[2].text = str(story_count)
        points = sum(
            s.get("story_points", 0)
            for f in features
            for s in f.get("user_stories", [])
        )
        row[3].text = str(points)

    # Determine output path
    if not output_path:
        output_dir = Path("output")
        output_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = str(output_dir / f"stories_{timestamp}.docx")
    else:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    doc.save(output_path)

    return {
        "file_path": output_path,
        "total_epics": len(epics),
        "total_stories": total_stories,
        "total_story_points": stories.get("total_story_points", 0),
        "message": f"Word document saved to: {output_path}",
    }


def _add_story_to_doc(doc, story: dict, story_num: int):
    """Add a single user story section to the Word document."""
    from docx.shared import Pt, RGBColor

    title = story.get("title", "Untitled Story")
    doc.add_heading(f"Story {story_num}: {title}", level=3)

    # Metadata line
    priority = story.get("priority", "Medium")
    points = story.get("story_points", 0)
    labels = ", ".join(story.get("labels", []))
    meta_text = f"Priority: {priority} | Story Points: {points}"
    if labels:
        meta_text += f" | Labels: {labels}"
    meta_para = doc.add_paragraph(meta_text)
    meta_para.runs[0].font.size = Pt(9)
    meta_para.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Description
    if story.get("description"):
        doc.add_paragraph(story["description"])

    # Acceptance Criteria
    acs = story.get("acceptance_criteria", [])
    if acs:
        ac_label = doc.add_paragraph()
        ac_run = ac_label.add_run("Acceptance Criteria:")
        ac_run.bold = True
        for ac in acs:
            doc.add_paragraph(ac, style="List Number")

    # Tasks
    tasks = story.get("tasks", [])
    if tasks:
        task_label = doc.add_paragraph()
        task_run = task_label.add_run("Tasks:")
        task_run.bold = True
        for task in tasks:
            doc.add_paragraph(task, style="List Bullet")

    doc.add_paragraph("")  # Spacing


@mcp.tool()
async def export_pdf(stories_json: str, output_path: str | None = None) -> dict:
    """
    Export generated stories to a PDF document.

    Creates a formatted PDF with epics, features, user stories,
    acceptance criteria, tasks, and a summary table.

    Args:
        stories_json: JSON string of generated stories (epics/features/stories structure).
        output_path: File path to save the .pdf file. If None, saves to output/stories_<timestamp>.pdf

    Returns:
        Dictionary with file path of the generated document and summary.
    """
    from pathlib import Path

    from fpdf import FPDF

    try:
        stories = json.loads(stories_json)
    except json.JSONDecodeError as e:
        return {"error": f"Invalid JSON: {e}"}

    epics = stories.get("epics", [])
    if not epics:
        return {"error": "No epics found in the stories JSON."}

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Title page
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 24)
    pdf.cell(0, 20, "Development Plan", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "User Stories", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(10)

    pdf.set_font("Helvetica", "", 11)
    client = state.connected_client or "Unknown Client"
    pdf.cell(0, 7, f"Client: {client}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 7, f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", new_x="LMARGIN", new_y="NEXT")
    pdf.cell(0, 7, f"Total Story Points: {stories.get('total_story_points', 'N/A')}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # Warnings
    warnings = stories.get("warnings", [])
    if warnings:
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, "Warnings", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(200, 0, 0)
        for w in warnings:
            pdf.multi_cell(0, 6, f"  * {w}", new_x="LMARGIN", new_y="NEXT")
        pdf.set_text_color(0, 0, 0)
        pdf.ln(3)

    # Assumptions
    assumptions = stories.get("assumptions", [])
    if assumptions:
        pdf.set_font("Helvetica", "B", 14)
        pdf.cell(0, 10, "Assumptions", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        for a in assumptions:
            pdf.multi_cell(0, 6, f"  * {a}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

    # Epics, Features, Stories
    total_stories = 0
    for epic_idx, epic in enumerate(epics, 1):
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 16)
        pdf.multi_cell(0, 9, f"Epic {epic_idx}: {epic.get('title', 'Untitled')}", new_x="LMARGIN", new_y="NEXT")
        if epic.get("description"):
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 6, epic["description"], new_x="LMARGIN", new_y="NEXT")
        pdf.ln(3)

        for feature in epic.get("features", []):
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 8, f"Feature: {feature.get('title', 'Untitled')}", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

            for story in feature.get("user_stories", []):
                total_stories += 1
                _add_story_to_pdf(pdf, story, total_stories)

    # Summary table
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Summary", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(3)

    # Table header
    pdf.set_font("Helvetica", "B", 10)
    col_widths = [70, 30, 30, 40]
    headers = ["Epic", "Features", "Stories", "Story Points"]
    for i, h in enumerate(headers):
        pdf.cell(col_widths[i], 8, h, border=1)
    pdf.ln()

    # Table rows
    pdf.set_font("Helvetica", "", 10)
    for epic in epics:
        features = epic.get("features", [])
        story_count = sum(len(f.get("user_stories", [])) for f in features)
        points = sum(
            s.get("story_points", 0)
            for f in features
            for s in f.get("user_stories", [])
        )
        pdf.cell(col_widths[0], 8, epic.get("title", "")[:35], border=1)
        pdf.cell(col_widths[1], 8, str(len(features)), border=1)
        pdf.cell(col_widths[2], 8, str(story_count), border=1)
        pdf.cell(col_widths[3], 8, str(points), border=1)
        pdf.ln()

    # Determine output path
    if not output_path:
        output_dir = Path("output")
        output_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = str(output_dir / f"stories_{timestamp}.pdf")
    else:
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)

    pdf.output(output_path)

    return {
        "file_path": output_path,
        "total_epics": len(epics),
        "total_stories": total_stories,
        "total_story_points": stories.get("total_story_points", 0),
        "message": f"PDF document saved to: {output_path}",
    }


def _add_story_to_pdf(pdf, story: dict, story_num: int):
    """Add a single user story section to the PDF document."""
    title = story.get("title", "Untitled Story")
    pdf.set_font("Helvetica", "B", 11)
    pdf.multi_cell(0, 7, f"Story {story_num}: {title}", new_x="LMARGIN", new_y="NEXT")

    # Metadata
    priority = story.get("priority", "Medium")
    points = story.get("story_points", 0)
    labels = ", ".join(story.get("labels", []))
    meta = f"Priority: {priority} | Story Points: {points}"
    if labels:
        meta += f" | Labels: {labels}"
    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 5, meta, new_x="LMARGIN", new_y="NEXT")
    pdf.set_text_color(0, 0, 0)

    # Description
    if story.get("description"):
        pdf.set_font("Helvetica", "", 10)
        pdf.multi_cell(0, 6, story["description"], new_x="LMARGIN", new_y="NEXT")

    # Acceptance Criteria
    acs = story.get("acceptance_criteria", [])
    if acs:
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 7, "Acceptance Criteria:", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        for i, ac in enumerate(acs, 1):
            pdf.multi_cell(0, 6, f"  {i}. {ac}", new_x="LMARGIN", new_y="NEXT")

    # Tasks
    tasks = story.get("tasks", [])
    if tasks:
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 7, "Tasks:", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        for task in tasks:
            pdf.multi_cell(0, 6, f"  * {task}", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(5)
